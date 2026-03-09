from flask import Flask, render_template, request, redirect, url_for, session, flash
import sqlite3
import os
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import tempfile
from docx import Document
import pdfkit
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import uuid
import os
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, 'instance', 'database.db')
from dotenv import load_dotenv
load_dotenv()





def apply_enhancements_to_resume(data, enhancements):
    """Применяет усиления к данным резюме"""
    
    # Копируем данные, чтобы не менять оригинал
    enhanced_data = data.copy()
    
    # 1. ATS-оптимизация - добавляем скрытые ключевые слова в навыки
    if 'ats' in enhancements:
        ats_keywords = {
            'python': 'Python, Django, Flask, SQL, Git, API, REST, Docker',
            'javascript': 'JavaScript, React, Node.js, HTML, CSS, jQuery',
            'менеджер': 'Управление проектами, Jira, Agile, Scrum, Kanban, Trello',
            'продажи': 'Продажи, Переговоры, CRM, Cold calls, В2B, В2С',
            'дизайнер': 'Figma, Adobe Photoshop, Illustrator, UI/UX, Prototyping',
            'маркетолог': 'SEO, SMM, Google Analytics, Яндекс.Метрика, Контекст'
        }
        
        # Ищем ключевые слова в должности
        position_lower = data['position'].lower()
        for key, keywords in ats_keywords.items():
            if key in position_lower:
                if enhanced_data['skills']:
                    enhanced_data['skills'] += ', ' + keywords
                else:
                    enhanced_data['skills'] = keywords
                break
    
    # 2. Помощь в написании - улучшаем описание
    if 'help' in enhancements:
        if not enhanced_data['about'] or len(enhanced_data['about']) < 50:
            templates = {
                'python': 'Опытный Python-разработчик с опытом создания веб-приложений. Участвовал в полном цикле разработки: от проектирования до поддержки. Использую современные подходы и лучшие практики.',
                'менеджер': 'Результативный менеджер проектов с опытом управления командами до 10 человек. Успешно завершил более 15 проектов в срок и в рамках бюджета.',
                'продажи': 'Активный и целеустремленный специалист по продажам. Выполняю и перевыполняю планы продаж. Умею находить подход к разным клиентам.',
                'дизайнер': 'Креативный дизайнер с портфолио успешных проектов. Создаю современный и функциональный дизайн, который решает задачи бизнеса.'
            }
            
            position_lower = data['position'].lower()
            for key, template in templates.items():
                if key in position_lower:
                    enhanced_data['about'] = template
                    break
    
    # 3. Проверка орфографии - исправляем ошибки
    if 'spell' in enhancements:
        # Словарь частых ошибок
        spelling = {
            'праграмист': 'программист',
            'праэкт': 'проект',
            'каманда': 'команда',
            'разроботка': 'разработка',
            'приложения': 'приложения',
            'оптимизация': 'оптимизация',
            'комуникабельный': 'коммуникабельный',
            'ответсвенный': 'ответственный',
            'обучаемость': 'обучаемость',
            'стрессоустойчивый': 'стрессоустойчивый'
        }
        
        # Проверяем все текстовые поля
        text_fields = ['about', 'description1']
        for field in text_fields:
            if enhanced_data.get(field):
                text = enhanced_data[field]
                for wrong, correct in spelling.items():
                    text = text.replace(wrong, correct)
                enhanced_data[field] = text
    
    return enhanced_data


app = Flask(__name__)
app.config['SECRET_KEY'] = 'vclab-secret-key-2024'

EMAIL_ADDRESS = 'rs1vcvsonya@gmail.com'
EMAIL_PASSWORD = os.environ.get('EMAIL_PASSWORD')  # Берем из переменной окружения
SMTP_SERVER = 'smtp.gmail.com'
SMTP_PORT = 587

# Функция для создания базы данных
def init_db():
    if not os.path.exists('instance'):
        os.makedirs('instance')
    
    conn = sqlite3.connect('DB_PATH')
    cur = conn.cursor()
    
    # Таблица резюме
    cur.execute('''
        CREATE TABLE IF NOT EXISTS resumes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            fullname TEXT NOT NULL,
            email TEXT NOT NULL,
            phone TEXT,
            city TEXT,
            position TEXT NOT NULL,
            about TEXT,
            company1 TEXT,
            position1 TEXT,
            period1 TEXT,
            description1 TEXT,
            university TEXT,
            graduation_year TEXT,
            specialty TEXT,
            skills TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Таблица пользователей
    cur.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            email TEXT NOT NULL UNIQUE,
            password TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    conn.commit()
    conn.close()
    print("✅ База данных создана!")

# Функция для создания PDF
def create_pdf(data, photo_path=None):
    try:
        # Передаем путь к фото в шаблон
        template_data = data.copy()
        if photo_path:
            abs_path = os.path.abspath(photo_path)
            template_data['photo'] = f'file:///{abs_path.replace("\\", "/")}'
        
        # Добавляем усиления в шаблон (если они есть)
        if 'enhancements' in data:
            template_data['enhancements'] = data['enhancements']
        
        # Рендерим HTML шаблон с данными
        html_string = render_template('resume/resume_template.html', **template_data)
        
        # Создаем временный файл для PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            # Генерируем PDF из HTML
            path_wkhtmltopdf = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
            config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
            pdfkit.from_string(html_string, tmp.name, configuration=config)
            return tmp.name
    except Exception as e:
        print(f"Ошибка создания PDF: {e}")
        # Если не получается создать PDF, создаем простой текстовый файл
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp.write(f"Резюме {data['fullname']}\n\nДолжность: {data['position']}".encode('utf-8'))
            return tmp.name

# Функция для создания Word документа
def create_word(data):
    doc = Document()
    
    # Заголовок
    title = doc.add_heading(data['fullname'], 0)
    title.alignment = 1  # Центрируем
    
    # Должность
    position = doc.add_heading(data['position'], level=1)
    position.alignment = 1
    
    # Контактная информация
    doc.add_paragraph()
    contact = doc.add_paragraph()
    contact.add_run(f"📧 {data['email']}").bold = True
    if data.get('phone'):
        contact.add_run(f" | 📞 {data['phone']}")
    if data.get('city'):
        contact.add_run(f" | 🏠 {data['city']}")
    
    # О себе
    if data.get('about'):
        doc.add_heading('О себе', level=2)
        doc.add_paragraph(data['about'])
    
    # Опыт работы
    if data.get('company1'):
        doc.add_heading('Опыт работы', level=2)
        p = doc.add_paragraph()
        p.add_run(data['company1']).bold = True
        if data.get('period1'):
            p.add_run(f" ({data['period1']})")
        if data.get('description1'):
            doc.add_paragraph(data['description1'], style='List Bullet')
    
    # Образование
    if data.get('university'):
        doc.add_heading('Образование', level=2)
        p = doc.add_paragraph()
        p.add_run(data['university']).bold = True
        if data.get('graduation_year'):
            p.add_run(f" - {data['graduation_year']}")
        if data.get('specialty'):
            doc.add_paragraph(data['specialty'])
    
    # Навыки
    if data.get('skills'):
        doc.add_heading('Навыки', level=2)
        skills = data['skills'].split(',')
        for skill in skills:
            doc.add_paragraph(skill.strip(), style='List Bullet')
    
    # Сохраняем во временный файл
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        return tmp.name

# Функция отправки на почту (обновленная)
def send_resume_email(recipient_email, data, photo_path=None):
    try:
        # Получаем выбранные усиления из данных
        enhancements = data.get('enhancements', [])
        
        # Применяем усиления к данным резюме
        enhanced_data = apply_enhancements_to_resume(data, enhancements)
        
        # Создаем файлы с усилениями
        pdf_path = create_pdf(enhanced_data, photo_path)
        word_path = create_word(enhanced_data, photo_path)
        
        # Создаем письмо
        msg = MIMEMultipart()
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = recipient_email
        msg['Subject'] = f"Ваше резюме - {data['fullname']}"
        
        # Формируем список выбранных усилений для письма
        enhancements_text = ""
        if enhancements:
            enhancements_text = "<h3>✨ Выбранные усиления:</h3><ul>"
            if 'photo' in enhancements:
                enhancements_text += "<li>📸 Фото в резюме</li>"
            if 'ats' in enhancements:
                enhancements_text += "<li>🤖 ATS-оптимизация (роботы пропустят)</li>"
            if 'cover' in enhancements:
                enhancements_text += "<li>✍️ Сопроводительное письмо</li>"
            if 'spell' in enhancements:
                enhancements_text += "<li>🔍 Проверка орфографии</li>"
            if 'help' in enhancements:
                enhancements_text += "<li>🎯 Помощь в написании</li>"
            enhancements_text += "</ul>"
        
        # HTML тело письма
        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; padding: 20px;">
            <h2 style="color: #2563eb;">Здравствуйте, {data['fullname']}!</h2>
            <p>Ваше резюме готово и прикреплено к этому письму.</p>
            
            <h3>📋 Краткая информация:</h3>
            <ul>
                <li><b>Должность:</b> {data['position']}</li>
                <li><b>Email:</b> {data['email']}</li>
                <li><b>Телефон:</b> {data.get('phone', 'не указан')}</li>
                {f'<li><b>Город:</b> {data["city"]}</li>' if data.get('city') else ''}
                {'<li><b>Фото:</b> добавлено</li>' if photo_path else ''}
            </ul>
            
            {enhancements_text}
            
            <p>Файлы прикреплены в двух форматах:</p>
            <ul>
                <li>📄 <b>PDF</b> - для печати и отправки</li>
                <li>📝 <b>Word</b> - для редактирования</li>
            </ul>
            
            <hr>
            <p style="color: #666; font-size: 12px;">
                Спасибо за использование VCLab!

                Если у вас есть вопросы, просто ответьте на это письмо.
            </p>
        </body>
        </html>
        """
        
        msg.attach(MIMEText(html_content, 'html'))
        
        # Прикрепляем PDF
        with open(pdf_path, 'rb') as f:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename="resume_{data["fullname"]}.pdf"'
            )
            msg.attach(part)
        
        # Прикрепляем Word
        with open(word_path, 'rb') as f:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename="resume_{data["fullname"]}.docx"'
            )
            msg.attach(part)
        
        # Отправляем через SMTP
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.send_message(msg)
        server.quit()
        
        # Удаляем временные файлы
        os.unlink(pdf_path)
        os.unlink(word_path)
        if photo_path and os.path.exists(photo_path):
            os.unlink(photo_path)
        
        print(f"✅ Письмо отправлено на {recipient_email}")
        return True
        
    except Exception as e:
        print(f"❌ Ошибка отправки: {e}")
        return False

# ========== МАРШРУТЫ (ROUTES) ==========

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/create')
def create():
    return render_template('create.html')

@app.route('/success')
def success():
    return render_template('success.html')

@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/dashboard')
def dashboard():
    # Проверяем, вошел ли пользователь
    if 'user_id' not in session:
        flash('Пожалуйста, войдите в систему', 'error')
        return redirect(url_for('login'))
    
    # Получаем резюме пользователя из базы
    conn = sqlite3.connect('DB_PATN')
    cur = conn.cursor()
    
    cur.execute('''
        SELECT * FROM resumes 
        WHERE user_id = ? 
        ORDER BY created_at DESC
    ''', (session['user_id'],))
    
    resumes = cur.fetchall()
    conn.close()
    
    return render_template('dashboard.html', resumes=resumes, username=session.get('username'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        
        conn = sqlite3.connect('DB_PATH')
        cur = conn.cursor()
        cur.execute('SELECT * FROM users WHERE email = ?', (email,))
        user = cur.fetchone()
        conn.close()
        
        if user and check_password_hash(user[3], password):
            session['user_id'] = user[0]
            session['username'] = user[1]
            flash('Вы успешно вошли!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Неверный email или пароль', 'error')
            
    return render_template('login.html')

@app.route('/register', methods=['POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        if password != confirm_password:
            flash('Пароли не совпадают', 'error')
            return redirect(url_for('login'))
        
        # Хешируем пароль
        hashed_password = generate_password_hash(password)
        
        try:
            conn = sqlite3.connect('DB_PATH')
            cur = conn.cursor()
            cur.execute('''
                INSERT INTO users (username, email, password)
                VALUES (?, ?, ?)
            ''', (username, email, hashed_password))
            conn.commit()
            conn.close()
            
            flash('Регистрация успешна! Теперь войдите', 'success')
        except sqlite3.IntegrityError:
            flash('Пользователь с таким email уже существует', 'error')
            
    return redirect(url_for('login'))

@app.route('/logout')
def logout():
    session.clear()
    flash('Вы вышли из системы', 'info')
    return redirect(url_for('index'))

@app.route('/create-resume', methods=['POST'])
def create_resume():
    if request.method == 'POST':
        # Получаем выбранные усиления
        enhancements = []
        total_price = 99  # Базовая цена
        
        if request.form.get('enhance_photo'):
            enhancements.append('photo')
            total_price += 99
        if request.form.get('enhance_ats'):
            enhancements.append('ats')
            total_price += 149
        if request.form.get('enhance_cover'):
            enhancements.append('cover')
            total_price += 149
        if request.form.get('enhance_spell'):
            enhancements.append('spell')
            total_price += 99
        if request.form.get('enhance_help'):
            enhancements.append('help')
            total_price += 199
        
        # Обработка фото
        photo_path = None
        if 'photo' in request.files and 'enhance_photo' in request.form:
            photo = request.files['photo']
            if photo and allowed_file(photo.filename):
                filename = secure_filename(photo.filename)
                ext = filename.rsplit('.', 1)[1].lower()
                new_filename = f"{uuid.uuid4()}.{ext}"
                photo_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
                photo.save(photo_path)
        
        # Собираем данные из формы
        data = {
            'fullname': request.form['fullname'],
            'email': request.form['email'],
            'phone': request.form.get('phone', ''),
            'city': request.form.get('city', ''),
            'position': request.form['position'],
            'about': request.form.get('about', ''),
            'company1': request.form.get('company1', ''),
            'position1': request.form.get('position1', ''),
            'period1': request.form.get('period1', ''),
            'description1': request.form.get('description1', ''),
            'university': request.form.get('university', ''),
            'graduation_year': request.form.get('graduation_year', ''),
            'specialty': request.form.get('specialty', ''),
            'skills': request.form.get('skills', ''),
            'enhancements': enhancements,
            'total_price': total_price
        }
        
        # ПРИМЕНЯЕМ УСИЛЕНИЯ
        
        # 1. Помощь в написании (добавляем подсказки, если поля пустые)
        if 'help' in enhancements:
            if not data['about']:
                data['about'] = get_about_template(data['position'])
            if not data['skills']:
                data['skills'] = get_skills_template(data['position'])
        
        # 2. Проверка орфографии (исправляем ошибки)
        if 'spell' in enhancements:
            data['about'] = check_spelling(data['about'])
            data['description1'] = check_spelling(data['description1'])
        
        # 3. ATS-оптимизация (добавляем ключевые слова)
        if 'ats' in enhancements:
            data['skills'] = add_ats_keywords(data['skills'], data['position'])
        
        # Сохраняем в базу
        conn = sqlite3.connect('DB_PATH')
        cur = conn.cursor()
        
        user_id = session.get('user_id', None)
        
        cur.execute('''
            INSERT INTO resumes (
                user_id, photo, fullname, email, phone, city, position, about,
                company1, position1, period1, description1,
                university, graduation_year, specialty, skills
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            user_id, photo_path, data['fullname'], data['email'], data['phone'], 
            data['city'], data['position'], data['about'], data['company1'], 
            data['position1'], data['period1'], data['description1'],
            data['university'], data['graduation_year'], 
            data['specialty'], data['skills']
        ))
        
        conn.commit()
        resume_id = cur.lastrowid
        conn.close()
        
        print(f"✅ Резюме #{resume_id} для {data['fullname']} сохранено в базу!")
        print(f"💰 Итоговая стоимость: {total_price} ₽")
        print(f"✨ Выбранные усиления: {enhancements}")
        
        # Отправляем на почту
        if send_resume_email(data['email'], data, photo_path):
            print(f"✅ Письмо успешно отправлено на {data['email']}")
            
            # Если выбрано сопроводительное письмо - отправляем отдельно
            if 'cover' in enhancements:
                send_cover_letter(data['email'], data)
        else:
            print(f"❌ Не удалось отправить письмо на {data['email']}")
        
        return redirect(url_for('success'))
@app.route('/delete-resume/<int:resume_id>', methods=['POST'])
def delete_resume(resume_id):
    if 'user_id' not in session:
        flash('Пожалуйста, войдите в систему', 'error')
        return redirect(url_for('login'))
    
    # Удаляем резюме из базы
    conn = sqlite3.connect('BD_PATH')
    cur = conn.cursor()
    
    # Проверяем, что резюме принадлежит пользователю
    cur.execute('DELETE FROM resumes WHERE id = ? AND user_id = ?', 
                (resume_id, session['user_id']))
    conn.commit()
    conn.close()
    
    flash('Резюме успешно удалено', 'success')
    return redirect(url_for('dashboard'))        
@app.route('/edit-resume/<int:resume_id>')
def edit_resume(resume_id):
    # Проверяем, вошел ли пользователь
    if 'user_id' not in session:
        flash('Пожалуйста, войдите в систему', 'error')
        return redirect(url_for('login'))
    
    # Получаем данные резюме из базы
    conn = sqlite3.connect('DB_PATH')
    cur = conn.cursor()
    
    cur.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', 
                (resume_id, session['user_id']))
    resume = cur.fetchone()
    conn.close()
    
    if not resume:
        flash('Резюме не найдено', 'error')
        return redirect(url_for('dashboard'))
    
    # Преобразуем кортеж в словарь для удобства
    resume_data = {
        'id': resume[0],
        'fullname': resume[2],
        'email': resume[3],
        'phone': resume[4],
        'city': resume[5],
        'position': resume[6],
        'about': resume[7],
        'company1': resume[8],
        'position1': resume[9],
        'period1': resume[10],
        'description1': resume[11],
        'university': resume[12],
        'graduation_year': resume[13],
        'specialty': resume[14],
        'skills': resume[15]
    }
    
    return render_template('edit_resume.html', resume=resume_data)

@app.route('/update-resume/<int:resume_id>', methods=['POST'])
def update_resume(resume_id):
    if 'user_id' not in session:
        flash('Пожалуйста, войдите в систему', 'error')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Собираем данные из формы
        data = {
            'fullname': request.form['fullname'],
            'email': request.form['email'],
            'phone': request.form.get('phone', ''),
            'city': request.form.get('city', ''),
            'position': request.form['position'],
            'about': request.form.get('about', ''),
            'company1': request.form.get('company1', ''),
            'position1': request.form.get('position1', ''),
            'period1': request.form.get('period1', ''),
            'description1': request.form.get('description1', ''),
            'university': request.form.get('university', ''),
            'graduation_year': request.form.get('graduation_year', ''),
            'specialty': request.form.get('specialty', ''),
            'skills': request.form.get('skills', '')
        }
        
        # Обновляем в базе
        conn = sqlite3.connect('DB_PATH')
        cur = conn.cursor()
        
        cur.execute('''
            UPDATE resumes SET
                fullname=?, email=?, phone=?, city=?, position=?, about=?,
                company1=?, position1=?, period1=?, description1=?,
                university=?, graduation_year=?, specialty=?, skills=?
            WHERE id=? AND user_id=?
        ''', (
            data['fullname'], data['email'], data['phone'], data['city'], 
            data['position'], data['about'], data['company1'], 
            data['position1'], data['period1'], data['description1'],
            data['university'], data['graduation_year'], 
            data['specialty'], data['skills'],
            resume_id, session['user_id']
        ))
        
        conn.commit()
        conn.close()
        
        flash('Резюме успешно обновлено!', 'success')
        
        # Отправляем обновленную версию на почту
        if send_resume_email(data['email'], data):
            flash('Обновленная версия отправлена на почту', 'success')
        
        return redirect(url_for('dashboard'))
@app.route('/view-resume/<int:resume_id>')
def view_resume(resume_id):
    # Проверяем, вошел ли пользователь
    if 'user_id' not in session:
        flash('Пожалуйста, войдите в систему', 'error')
        return redirect(url_for('login'))
    
    # Получаем данные резюме из базы
    conn = sqlite3.connect('DB_PATH')
    cur = conn.cursor()
    
    cur.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', 
                (resume_id, session['user_id']))
    resume = cur.fetchone()
    conn.close()
    
    if not resume:
        flash('Резюме не найдено', 'error')
        return redirect(url_for('dashboard'))
    
    # Преобразуем кортеж в словарь
    resume_data = {
        'id': resume[0],
        'fullname': resume[2],
        'email': resume[3],
        'phone': resume[4],
        'city': resume[5],
        'position': resume[6],
        'about': resume[7],
        'company1': resume[8],
        'position1': resume[9],
        'period1': resume[10],
        'description1': resume[11],
        'university': resume[12],
        'graduation_year': resume[13],
        'specialty': resume[14],
        'skills': resume[15]
    }
    
    return render_template('view_resume.html', resume=resume_data)

@app.route('/download-resume/<int:resume_id>/<string:format>')
def download_resume(resume_id, format):
    # Проверяем, вошел ли пользователь
    if 'user_id' not in session:
        flash('Пожалуйста, войдите в систему', 'error')
        return redirect(url_for('login'))
    
    # Получаем данные резюме из базы
    conn = sqlite3.connect('DB_PATH')
    cur = conn.cursor()
    
    cur.execute('SELECT * FROM resumes WHERE id = ? AND user_id = ?', 
                (resume_id, session['user_id']))
    resume = cur.fetchone()
    conn.close()
    
    if not resume:
        flash('Резюме не найдено', 'error')
        return redirect(url_for('dashboard'))
    
    # Преобразуем кортеж в словарь
    resume_data = {
        'fullname': resume[2],
        'email': resume[3],
        'phone': resume[4],
        'city': resume[5],
        'position': resume[6],
        'about': resume[7],
        'company1': resume[8],
        'position1': resume[9],
        'period1': resume[10],
        'description1': resume[11],
        'university': resume[12],
        'graduation_year': resume[13],
        'specialty': resume[14],
        'skills': resume[15]
    }
    
    # Создаем файл в нужном формате
    if format == 'pdf':
        file_path = create_pdf(resume_data)
        mimetype = 'application/pdf'
        filename = f"resume_{resume_data['fullname']}.pdf"
    elif format == 'word':
        file_path = create_word(resume_data)
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        filename = f"resume_{resume_data['fullname']}.docx"
    else:
        flash('Неверный формат файла', 'error')
        return redirect(url_for('dashboard'))
    
    # Отправляем файл пользователю
    from flask import send_file
    return send_file(file_path, mimetype=mimetype, as_attachment=True, download_name=filename)
@app.route('/faq')
def faq():
    return render_template('faq.html')
# Создаем базу данных при запуске
with app.app_context():
    init_db()
    print("✅ База данных создана/проверена")
if __name__ == '__main__':
    init_db()
    app.run(debug=True)